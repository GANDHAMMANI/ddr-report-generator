import base64
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.models.pipeline import DDRReport, SeverityLevel
from app.core.image_handler import get_severity_color
from app.utils.logger import logger


SEVERITY_COLORS = {
    "Critical": RGBColor(0xDC, 0x26, 0x26),
    "High": RGBColor(0xEA, 0x58, 0x0C),
    "Medium": RGBColor(0xD9, 0x77, 0x06),
    "Low": RGBColor(0x16, 0xA3, 0x4A),
    "Not Available": RGBColor(0x6B, 0x72, 0x80),
}

PRIORITY_COLORS = {
    "Immediate": RGBColor(0xDC, 0x26, 0x26),
    "Short-term": RGBColor(0xD9, 0x77, 0x06),
    "Long-term": RGBColor(0x25, 0x63, 0xEB),
}

HEADER_COLOR = RGBColor(0x1E, 0x3A, 0x5F)


def _add_section_heading(doc: Document, number: int, title: str):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = HEADER_COLOR
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = HEADER_COLOR

    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc: Document, text: str, color: RGBColor = None):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color


def _add_image_to_doc(doc: Document, image_data, area_name: str):
    """Add an image to the DOCX document, converting via Pillow for compatibility."""
    from PIL import Image as PILImage
    try:
        img_bytes = base64.b64decode(image_data.base64_data)
        # Convert through Pillow → ensures valid JPEG/PNG for python-docx
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        # Convert to RGB if needed (handles RGBA, P, CMYK etc)
        if pil_img.mode not in ("RGB", "L"):
            pil_img = pil_img.convert("RGB")
        output = io.BytesIO()
        pil_img.save(output, format="JPEG", quality=85)
        output.seek(0)
        caption = image_data.caption or f"{image_data.source.title()} Image"
        doc.add_picture(output, width=Inches(4.5))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        p.runs[0].italic = True
        logger.info(f"Added image for {area_name}")
    except Exception as e:
        logger.warning(f"Could not add image for {area_name}: {e} | b64 length: {len(image_data.base64_data) if image_data.base64_data else 0}")
        p = doc.add_paragraph("📷 Image Not Available")
        p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        p.runs[0].italic = True


def export_docx(ddr: DDRReport, output_dir: Path) -> Path:
    """Export DDR report as a professionally formatted DOCX file."""
    logger.info("Exporting DDR to DOCX...")
    doc = Document()

    # ── Page Setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Title Block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DETAILED DIAGNOSTIC REPORT (DDR)")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = HEADER_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("AI-Generated Professional Property Assessment")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub_run.italic = True

    doc.add_paragraph()

    # ── Property Info Table ───────────────────────────────────────────────────
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    cells = [
        ("Property", ddr.property_name),
        ("Address", ddr.property_address),
        ("Inspection Date", ddr.inspection_date),
        ("Report Generated", ddr.generated_at),
    ]
    for i, (label, value) in enumerate(cells):
        row = table.rows[i // 2]
        cell = row.cells[i % 2]
        cell.text = ""
        p = cell.paragraphs[0]
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.color.rgb = HEADER_COLOR
        p.add_run(value or "Not Available")

    doc.add_paragraph()

    # ── Section 1: Summary ────────────────────────────────────────────────────
    _add_section_heading(doc, 1, "Property Issue Summary")
    p = doc.add_paragraph(ddr.section_1_summary)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

    # ── Section 2: Area-wise ──────────────────────────────────────────────────
    _add_section_heading(doc, 2, "Area-wise Observations")
    for area in ddr.section_2_area_wise:
        # Area subheading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(f"📍 {area.area_name}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = HEADER_COLOR

        if area.observations:
            p = doc.add_paragraph()
            run = p.add_run("Visual Inspection:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            for obs in area.observations:
                _add_bullet(doc, obs)

        if area.thermal_findings:
            p = doc.add_paragraph()
            run = p.add_run("🌡️ Thermal Findings:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            for tf in area.thermal_findings:
                _add_bullet(doc, tf, RGBColor(0x1E, 0x40, 0xAF))

        if area.has_conflict and area.conflict_note:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ Conflict: {area.conflict_note}")
            run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
            run.italic = True

        # Images — max 3 per area to keep file size manageable
        if area.images:
            for img in area.images[:3]:
                _add_image_to_doc(doc, img, area.area_name)
        else:
            p = doc.add_paragraph("📷 Image Not Available")
            p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            p.runs[0].italic = True

    # ── Section 3: Root Cause ─────────────────────────────────────────────────
    _add_section_heading(doc, 3, "Probable Root Cause")
    p = doc.add_paragraph(ddr.section_3_root_cause)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Section 4: Severity ───────────────────────────────────────────────────
    _add_section_heading(doc, 4, "Severity Assessment")
    for item in ddr.section_4_severity:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        color = SEVERITY_COLORS.get(item.severity, RGBColor(0x6B, 0x72, 0x80))
        run = p.add_run(f"[{item.severity.upper()}] ")
        run.bold = True
        run.font.color.rgb = color
        run2 = p.add_run(f"{item.area_name}: ")
        run2.bold = True
        p.add_run(item.reasoning)

    # ── Section 5: Actions ────────────────────────────────────────────────────
    _add_section_heading(doc, 5, "Recommended Actions")
    for action in ddr.section_5_actions:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        color = PRIORITY_COLORS.get(action.priority, RGBColor(0x6B, 0x72, 0x80))
        run = p.add_run(f"[{action.priority.upper()}] ")
        run.bold = True
        run.font.color.rgb = color
        run2 = p.add_run(f"{action.area_name} — ")
        run2.bold = True
        p.add_run(f"{action.action} ({action.timeline})")

    # ── Section 6: Notes ──────────────────────────────────────────────────────
    _add_section_heading(doc, 6, "Additional Notes")
    p = doc.add_paragraph(ddr.section_6_notes)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Section 7: Missing Info ───────────────────────────────────────────────
    _add_section_heading(doc, 7, "Missing or Unclear Information")
    if ddr.section_7_missing:
        for item in ddr.section_7_missing:
            _add_bullet(doc, item, RGBColor(0x9C, 0xA3, 0xAF))
    else:
        doc.add_paragraph("No missing information identified.")

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(f"Report ID: {ddr.report_id}  |  Generated: {ddr.generated_at}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    output_path = output_dir / "ddr_report.docx"
    doc.save(str(output_path))
    logger.info(f"DOCX exported → {output_path}")
    return output_path